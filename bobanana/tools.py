import os

"""Agent 工具集 — 文档解析、网络搜索、嵌入模型、文本分块。"""

import concurrent.futures
import logging
import re
import threading
import time
from pathlib import Path

from bobanana.config import (
    BASE_DIR,
    EMBEDDING_PROVIDER,
    OPENAI_EMBEDDING_MODEL,
    SENTENCE_TRANSFORMERS_MODEL,
)

logger = logging.getLogger(__name__)

# LLM 懒加载缓存（P0 修复：避免 get_llm 首次调用时 NameError）
_llm = None
_llm_cache: dict = {}              # provider -> LLM 实例(可能为 None)
_provider_circuit: dict = {}       # provider -> 冷却截止 monotonic 时间戳
_provider_circuit_lock = threading.Lock()
_PROVIDER_COOLDOWN_SEC = 60.0
_llm_executor = None
_llm_executor_lock = threading.Lock()
class _ExecutorProxy:
    """共享线程池代理：submit 转发给真实池，shutdown 为空操作，避免误关共享池。"""

    def __init__(self, executor: concurrent.futures.ThreadPoolExecutor):
        self._executor = executor

    def submit(self, *args, **kwargs):
        return self._executor.submit(*args, **kwargs)

    def shutdown(self, *args, **kwargs):
        return None


def _get_llm_executor() -> _ExecutorProxy:
    """复用全局线程池，避免每次 LLM 调用都创建/销毁 executor。"""
    global _llm_executor
    if _llm_executor is None:
        with _llm_executor_lock:
            if _llm_executor is None:
                _llm_executor = concurrent.futures.ThreadPoolExecutor(
                    max_workers=8, thread_name_prefix="study-wiki-llm"
                )
    return _ExecutorProxy(_llm_executor)
def reset_llm_cache() -> None:
    """清空 LLM 实例缓存与熔断状态，配置变更后调用。"""
    global _llm
    _llm = None
    _llm_cache.clear()
    with _provider_circuit_lock:
        _provider_circuit.clear()


# ═══════════════════════════════════════════════════════════
# 1. 嵌入模型
# ═══════════════════════════════════════════════════════════

_embedding_model = None
_embedding_lock = threading.Lock()

# 打包进安装包的内嵌嵌入模型目录(CI 在构建时下载 all-MiniLM-L6-v2 放入)。
# 存在则完全离线加载;不存在则回退到 HF 缓存/联网下载。
_BUNDLED_EMBEDDING_DIR = BASE_DIR / "vendor_model"


def _bundled_embedding_ready() -> bool:
    return (
        (_BUNDLED_EMBEDDING_DIR / "model.safetensors").exists()
        or (_BUNDLED_EMBEDDING_DIR / "pytorch_model.bin").exists()
    ) and (_BUNDLED_EMBEDDING_DIR / "config.json").exists()


def _resolve_cached_model_path(model_name: str) -> str | None:
    """解析本地 HF 缓存快照路径,按本地目录加载以绕开 hub 网络检查。

    sentence_transformers 即使设置 HF_HUB_OFFLINE=1 仍会对缺失的
    adapter_config.json 发 HEAD 请求并重试 5 次,网络不可达时每次卡数十秒。
    直接解析缓存目录可完全离线、秒级加载。
    """
    try:
        cache_root = Path(
            os.environ.get("HF_HOME")
            or os.path.expanduser("~/.cache/huggingface/hub")
        )
        snap_dir = cache_root / ("models--" + model_name.replace("/", "--")) / "snapshots"
        if not snap_dir.is_dir():
            return None
        for snap in sorted(snap_dir.iterdir(), reverse=True):
            if snap.is_dir() and (snap / "config.json").exists():
                return str(snap)
    except Exception:
        return None
    return None


def _block_hf_network() -> None:
    """把 HF hub 端点指到本地黑洞,使任何 hub 请求立即失败(连接拒绝)。

    sentence_transformers 加载时会无视离线标志做 adapter_config.json 的
    HEAD 探测;在离线/受限网络环境,该探测的重试(指数退避 + 连接超时)
    会把一次模型加载拖到数分钟。黑洞端点使其瞬间失败,离线源照常加载。
    """
    os.environ["HF_HUB_OFFLINE"] = "1"
    os.environ["TRANSFORMERS_OFFLINE"] = "1"
    try:
        import huggingface_hub.constants as _hub_constants
        # ENDPOINT 与 URL 模板都是导入时拼接的常量,必须同时改,
        # 否则文件 HEAD 仍会打到真实的 huggingface.co
        _hub_constants.ENDPOINT = "http://127.0.0.1:9"
        _hub_constants.HUGGINGFACE_CO_URL_TEMPLATE = (
            "http://127.0.0.1:9/{repo_id}/resolve/{revision}/{filename}"
        )
    except Exception:  # noqa: BLE001 — 兜底: 失败也不影响加载
        pass
    try:
        # 关闭 hub 请求的指数退避重试(默认 5 次,离线时每次白等数十秒)
        import huggingface_hub.utils._http as _hh
        if not getattr(_hh, "_studywiki_patched", False):
            _orig_backoff = _hh.http_backoff

            def _no_retry(*args, **kwargs):
                kwargs["max_retries"] = 0
                kwargs["base_wait_time"] = 0
                kwargs["max_wait_time"] = 0
                return _orig_backoff(*args, **kwargs)

            _hh.http_backoff = _no_retry
            _hh._studywiki_patched = True
            try:
                import huggingface_hub.hf_api as _hfa
                if _hfa.http_backoff is _orig_backoff:
                    _hfa.http_backoff = _no_retry
            except Exception:  # noqa: BLE001
                pass
    except Exception:  # noqa: BLE001 — 兜底: 失败也不影响加载
        pass


def get_embedding_model():
    """懒加载嵌入模型。

    加载顺序: 内嵌模型目录(离线) → 本地 HF 缓存 → 联网下载。
    仅在确定可用离线源时才强制 HF 离线模式,避免新机器上因无缓存而加载失败。
    """
    global _embedding_model
    if _embedding_model is not None:
        return _embedding_model

    with _embedding_lock:
        if _embedding_model is not None:
            return _embedding_model

    if EMBEDDING_PROVIDER == "sentence-transformers":
        from sentence_transformers import SentenceTransformer
        # 统一在网络黑洞下加载: 本地/内嵌模型不受影响,任何 hub 探测立即失败
        _block_hf_network()
        if _bundled_embedding_ready():
            model_ref = str(_BUNDLED_EMBEDDING_DIR)
            os.environ["HF_HUB_OFFLINE"] = "1"
            os.environ["TRANSFORMERS_OFFLINE"] = "1"
            logger.info("加载内嵌嵌入模型: %s ...", model_ref)
            try:
                _embedding_model = SentenceTransformer(model_ref)
            except Exception as e:
                raise RuntimeError(
                    f"内嵌嵌入模型加载失败({model_ref}): {e}"
                ) from e
        else:
            model_ref = SENTENCE_TRANSFORMERS_MODEL
            os.environ["HF_HUB_OFFLINE"] = "1"
            os.environ["TRANSFORMERS_OFFLINE"] = "1"
            # 优先解析本地缓存目录直接加载(绕开 hub 网络重试,秒级)
            cached_path = _resolve_cached_model_path(model_ref)
            if cached_path:
                logger.info("加载嵌入模型(本地缓存 %s) ...", cached_path)
                try:
                    _embedding_model = SentenceTransformer(cached_path)
                except Exception as e:
                    raise RuntimeError(
                        f"嵌入模型加载失败({cached_path}): {e}"
                    ) from e
            else:
                logger.info("本地缓存缺失,尝试标准加载(可能联网) ...")
                try:
                    _embedding_model = SentenceTransformer(model_ref)
                except Exception:
                    if os.environ.get("STUDYWIKI_TEST_MODE") == "1":
                        raise
                    logger.warning("缓存缺失,尝试联网下载(首次约 90MB) ...")
                    os.environ.pop("HF_HUB_OFFLINE", None)
                    os.environ.pop("TRANSFORMERS_OFFLINE", None)
                    try:
                        _embedding_model = SentenceTransformer(model_ref)
                    except Exception as e:
                        raise RuntimeError(
                            f"嵌入模型加载失败({model_ref}): {e}。"
                            "请检查网络后重启,或使用包含 vendor_model/ 的安装包。"
                        ) from e
        logger.info("嵌入模型就绪, 维度: %d", _embedding_model.get_embedding_dimension())
    else:
        # OpenAI embedding: 直接返回 None, 由 embed_text 处理
        _embedding_model = "openai"
        logger.info("使用 OpenAI 嵌入: %s", OPENAI_EMBEDDING_MODEL)

    return _embedding_model

def embed_text(text: str) -> list[float]:
    """将文本转为向量。"""
    model = get_embedding_model()

    if model == "openai":
        from langchain_openai import OpenAIEmbeddings

        from bobanana.config import OPENAI_API_KEY
        emb = OpenAIEmbeddings(
            model=OPENAI_EMBEDDING_MODEL,
            openai_api_key=OPENAI_API_KEY,
        )
        vector = emb.embed_query(text)
        return vector
    else:
        vector = model.encode(text, normalize_embeddings=True).tolist()
        return vector

# ═══════════════════════════════════════════════════════════
# 2. 文档解析
# ═══════════════════════════════════════════════════════════

def parse_document(file_path: str, progress_callback=None) -> list[dict]:
    """解析文档，返回 [{page_num, text}, ...] 列表。"""
    path = Path(file_path)
    ext = path.suffix.lower()

    logger.info("解析文档: %s", file_path)

    if ext == ".pdf":
        return _parse_pdf(file_path, progress_callback)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_path)
    elif ext == ".md":
        return _parse_markdown(file_path)
    elif ext == ".txt":
        return _parse_text(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")

def _parse_pdf(file_path: str, progress_callback=None) -> list[dict]:
    """解析 PDF — 逐页提取文本，文本少于 20 字时尝试 OCR。单页超时 30s。"""
    import concurrent.futures

    import fitz
    pages: list[dict] = []
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        logger.error("PDF 打开失败，尝试图片 OCR: %s", e)
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
            return [{"page_num": 1, "text": text.strip()}]
        except Exception:
            return [{"page_num": 1, "text": ""}]
    total = len(doc)
    logger.info("PDF 共 %d 页", total)
    import threading
    result = {"pages": pages, "done": False}
    def _parse_all():
        for page_num in range(total):
            if result["done"]:
                break
            try:
                page = doc[page_num]
                with concurrent.futures.ThreadPoolExecutor(1) as pool:
                    f = pool.submit(page.get_text)
                    try:
                        text = f.result(timeout=30).strip()
                    except concurrent.futures.TimeoutError:
                        logger.warning("第 %d 页解析超时", page_num + 1)
                        text = ""
                if len(text) < 20:
                    ocr_text = _ocr_page(page)
                    if ocr_text and len(ocr_text) > len(text):
                        text = ocr_text
                # 跳过完全空页（OCR 后仍为空）
                if text.strip():
                    pages.append({"page_num": page_num + 1, "text": text})
            except Exception as e:
                logger.warning("第 %d 页异常: %s", page_num + 1, e)
                # 异常时跳过该页
            if progress_callback and (page_num + 1) % 5 == 0:
                try:
                    progress_callback({"stage": "parse", "current": page_num + 1, "total": total})
                except Exception:
                    pass
        doc.close()
        result["done"] = True

    # 整个 PDF 解析总超时 300s
    t = threading.Thread(target=_parse_all, daemon=True)
    t.start()
    t.join(timeout=300)
    result["done"] = True
    if t.is_alive():
        logger.error("PDF 解析总超时 (300s): %s", file_path)
    logger.info("PDF 解析完成: %d 页 (总 %d 页)", len(pages), total)
    return pages

def _ocr_page(page) -> str:
    """对 PyMuPDF 页面做 OCR。失败时返回空字符串。"""
    try:
        # 设置 tesseract 路径（不在 PATH 时的 fallback）
        import os as _os

        import pytesseract
        for _p in [r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                   r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe']:
            if _os.path.exists(_p):
                pytesseract.pytesseract.tesseract_cmd = _p
                break
        import io

        from PIL import Image
        pix = page.get_pixmap(dpi=300)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(img, lang="chi_sim+eng").strip()
        if text:
            return text
    except Exception as e:
        logger.debug("OCR 失败 (页 %d): %s", page.number + 1, e)
    return ""

def _parse_docx(file_path: str) -> list[dict]:
    """解析 Word — 按段落分页。"""
    from docx import Document
    pages = []
    doc = Document(file_path)
    current_text: list[str] = []
    page_num = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 分页符检测
        if "PAGE_BREAK" in text or "—————" in text:
            if current_text:
                pages.append({"page_num": page_num, "text": "\n".join(current_text)})
                page_num += 1
                current_text = []
        else:
            current_text.append(text)

    if current_text:
        pages.append({"page_num": page_num, "text": "\n".join(current_text)})

    # 如果没有分页符，整个文档作为一页
    if not pages:
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        pages.append({"page_num": 1, "text": full_text})

    logger.info("Word 解析完成: %d 页", len(pages))
    return pages

def _parse_markdown(file_path: str) -> list[dict]:
    """解析 Markdown — 按标题分节。"""
    import markdown as md_lib
    with open(file_path, encoding="utf-8") as f:
        raw = f.read()

    # 按 ## 标题分割
    sections = re.split(r"\n(?=##\s)", raw)
    pages = []
    for i, section in enumerate(sections):
        if section.strip():
            # 转 HTML 用于展示，但保留纯文本用于提取
            html = md_lib.markdown(section)
            pages.append({"page_num": i + 1, "text": section.strip(), "html": html})

    if not pages:
        pages.append({"page_num": 1, "text": raw})

    logger.info("Markdown 解析完成: %d 节", len(pages))
    return pages

def _parse_text(file_path: str) -> list[dict]:
    """解析纯文本 — 按空行分块。"""
    with open(file_path, encoding="utf-8") as f:
        text = f.read()

    blocks = re.split(r"\n\s*\n", text)
    pages = []
    for i, block in enumerate(blocks):
        if block.strip():
            pages.append({"page_num": i + 1, "text": block.strip()})

    if not pages:
        pages.append({"page_num": 1, "text": text})

    logger.info("文本解析完成: %d 块", len(pages))
    return pages

# ═══════════════════════════════════════════════════════════
# 3. 文本分块
# ═══════════════════════════════════════════════════════════

def chunk_text(text: str, max_chars: int = 500, overlap: int = 0) -> list[str]:
    """将长文本按最大字符数分块，可配置重叠字符数。"""
    if not text:
        return [""] if text == "" else []
    if len(text) <= max_chars:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + max_chars
        if end >= len(text):
            chunks.append(text[start:])
            break
        # 尽量在段落或句子边界处切断
        cut = text.rfind("\n\n", start, end)
        if cut <= start:
            cut = text.rfind("\n", start, end)
        if cut <= start:
            cut = text.rfind(". ", start, end)
        if cut <= start:
            cut = text.rfind(" ", start, end)
        if cut <= start:
            cut = end
        else:
            cut += 1  # 包含分隔符
        chunks.append(text[start:cut])
        start = cut - overlap
    return chunks


def web_search(query: str, top_k: int = 3) -> list[dict]:
    """使用 DuckDuckGo 搜索，返回 [{title, snippet, url}]。"""
    try:
        from duckduckgo_search import DDGS
        logger.info("网络搜索: %s", query)
        results = []
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=top_k):
                results.append({
                    "title": r.get("title", ""),
                    "snippet": r.get("body", ""),
                    "url": r.get("href", ""),
                })
        logger.info("搜索结果: %d 条", len(results))
        return results
    except Exception as e:
        logger.warning("网络搜索失败: %s", e)
        return []

# ═══════════════════════════════════════════════════════════
# 5. 知识提取辅助
# ═══════════════════════════════════════════════════════════

def _providers() -> list[str]:
    """返回降级链 provider 列表(按 LLM_PROVIDERS 顺序)。"""
    from bobanana.config import LLM_PROVIDER, LLM_PROVIDERS
    raw = (LLM_PROVIDERS or "").strip()
    if not raw:
        raw = (LLM_PROVIDER or "deepseek").strip()
    providers = [p.strip().lower() for p in raw.split(",") if p.strip()]
    return providers or ["deepseek"]


def _construct_llm(provider: str):
    """按 provider 名构造 LLM 实例。构造失败/无凭据时返回 None。"""
    from bobanana.config import (
        DEEPSEEK_API_KEY,
        DEEPSEEK_BASE_URL,
        DEEPSEEK_MODEL,
        LLM_TEMPERATURE,
        OLLAMA_BASE_URL,
        OLLAMA_MODEL,
        OPENAI_API_KEY,
        OPENAI_BASE_URL,
        OPENAI_MODEL,
    )

    provider = (provider or "").strip().lower()
    if provider == "deepseek":
        from langchain_openai import ChatOpenAI
        api_key = DEEPSEEK_API_KEY or OPENAI_API_KEY
        if not api_key:
            logger.debug("DeepSeek 无 API Key, 跳过")
            return None
        return ChatOpenAI(
            model=DEEPSEEK_MODEL,
            temperature=LLM_TEMPERATURE,
            api_key=api_key,
            base_url=DEEPSEEK_BASE_URL,
        )
    if provider == "openai":
        from langchain_openai import ChatOpenAI
        if not OPENAI_API_KEY:
            logger.debug("OpenAI 无 API Key, 跳过")
            return None
        return ChatOpenAI(
            model=OPENAI_MODEL,
            temperature=LLM_TEMPERATURE,
            api_key=OPENAI_API_KEY,
            base_url=OPENAI_BASE_URL,
        )
    if provider == "ollama":
        try:
            from langchain_community.chat_models import ChatOllama
            return ChatOllama(
                model=OLLAMA_MODEL,
                base_url=OLLAMA_BASE_URL,
                temperature=LLM_TEMPERATURE,
            )
        except Exception:
            try:
                from langchain_community.llms.ollama import Ollama
                return Ollama(
                    model=OLLAMA_MODEL,
                    base_url=OLLAMA_BASE_URL,
                    temperature=LLM_TEMPERATURE,
                )
            except Exception:
                logger.debug("Ollama 集成不可用, 跳过")
                return None
    return None


def _is_circuit_open(provider: str) -> bool:
    """检查 provider 是否处于熔断冷却期。"""
    with _provider_circuit_lock:
        until = _provider_circuit.get(provider)
        if until is None:
            return False
        if time.monotonic() < until:
            return True
        _provider_circuit.pop(provider, None)
        return False


def _trip_circuit(provider: str) -> None:
    """熔断 provider, 冷却 60s。"""
    with _provider_circuit_lock:
        _provider_circuit[provider] = time.monotonic() + _PROVIDER_COOLDOWN_SEC


def get_llm(provider: str = None):
    """懒加载 LLM 实例。

    provider 缺省时按降级链顺序返回第一个可用(未熔断且可构造)的候选,
    并更新 _llm 以保持旧调用兼容。
    """
    global _llm
    if provider is not None:
        p = provider.strip().lower()
        if p not in _llm_cache:
            inst = _construct_llm(p)
            _llm_cache[p] = inst
            if inst is not None:
                _llm = inst
                logger.info("LLM 就绪: %s", p)
        return _llm_cache.get(p)

    for p in _providers():
        if _is_circuit_open(p):
            continue
        inst = get_llm(p)
        if inst is not None:
            return inst

    # 全部熔断或不可构造时, 返回第一个可构造实例(即便熔断, 供调用方决策)
    for p in _providers():
        inst = get_llm(p)
        if inst is not None:
            return inst
    return None

# ═══════════════════════════════════════════════════════════
# 7. 文档预扫描器 (Phase 1)
# ═══════════════════════════════════════════════════════════

class ScanResult:
    """预扫描结果。"""
    def __init__(
        self, total_pages=0, valid_ranges=None, language="zh",
        doc_type="unknown", skipped_pages=None, pages=None,
    ):
        self.total_pages = total_pages
        self.valid_ranges = valid_ranges or []  # [(start, end, topic), ...]
        self.language = language
        self.doc_type = doc_type
        self.skipped_pages = skipped_pages or []
        self.pages = pages or []

class DocumentScanner:
    """Phase 1: 预扫描文档结构，识别有效内容区间。"""

    MIN_CONTENT_CHARS = 50  # 少于 50 个字符的页视为空白页

    def scan(self, file_path: str, progress_callback=None) -> ScanResult:
        pages = parse_document(file_path, progress_callback=progress_callback)
        if not pages:
            return ScanResult()

        total = len(pages)
        stats = self._analyze_pages(pages)
        structure = self._detect_structure(pages[:min(3, total)])
        valid_ranges, skipped = self._compute_valid_ranges(pages, stats)

        logger.info(
            "预扫描: %d 页, 有效区间 %d 个, 跳过 %d 页, 类型=%s, 语言=%s",
            total, len(valid_ranges), len(skipped),
            structure.get("doc_type", "?"), structure.get("language", "?"),
        )
        return ScanResult(
            total_pages=total,
            valid_ranges=valid_ranges,
            language=structure.get("language", "zh"),
            doc_type=structure.get("doc_type", "unknown"),
            skipped_pages=skipped,
            pages=pages,
        )

    def _analyze_pages(self, pages: list) -> list:
        """统计每页特征。"""
        stats = []
        for i, p in enumerate(pages):
            text = p.get("text", "")
            char_count = len(text.strip())
            non_space = len(text.strip().replace(" ", "").replace("\n", ""))
            stats.append({
                "index": i,
                "page_num": p.get("page_num", i + 1),
                "char_count": char_count,
                "non_space_chars": non_space,
                "is_blank": non_space < self.MIN_CONTENT_CHARS,
                "has_chinese": bool(re.search(r'[\u4e00-\u9fff]', text)),
            })
        return stats

    def _detect_structure(self, sample_pages: list) -> dict:
        """LLM 快速识别文档类型和语言。"""
        sample_text = "\n".join([p.get("text", "")[:500] for p in sample_pages if p.get("text")])[:2000]
        if not sample_text.strip():
            return {"language": "unknown", "doc_type": "unknown"}

        try:
            prompt = f"""分析以下文档开头的内容，返回 JSON:
{{
  "language": "zh" 或 "en",
  "doc_type": "教材" | "论文" | "PPT讲义" | "实验报告" | "其他",
  "title_hint": "可能的标题或主题"
}}

内容:
{sample_text[:1000]}"""
            result = llm_invoke("你是一个文档分析专家。只返回 JSON。", prompt)
            import json as _json
            for line in result.split("\n"):
                line = line.strip()
                if line.startswith("{"):
                    return _json.loads(line)
        except Exception:
            pass
        return {"language": "zh", "doc_type": "unknown"}

    def _compute_valid_ranges(self, pages: list, stats: list) -> tuple:
        """跳过空白页/封面/引用页，合并连续有效页为区间。"""
        valid_ranges = []
        skipped = []
        i = 0
        while i < len(pages):
            if stats[i]["is_blank"]:
                skipped.append(stats[i]["page_num"])
                i += 1
                continue

            # 检查是否可能是封面/目录（文档较长时才启用，避免误判短文档）
            if len(pages) > 3 and i < 3 and stats[i]["non_space_chars"] < 200:
                skipped.append(stats[i]["page_num"])
                i += 1
                continue

            # 合并连续有效页（每段最多 10 页，避免 LLM 响应过长被截断）
            start = stats[i]["page_num"]
            page_count = 0
            while i < len(pages) and not stats[i]["is_blank"] and page_count < 10:
                page_count += 1
                i += 1
            end = stats[i - 1]["page_num"]
            valid_ranges.append((start, end, f"第{start}-{end}页"))

        return valid_ranges, skipped

def _is_chat_model(llm) -> bool:
    """判断 LLM 实例是否为 Chat 模型(接受 messages 而非字符串)。"""
    try:
        from langchain_core.language_models.chat_models import BaseChatModel
        return isinstance(llm, BaseChatModel)
    except Exception:
        return True


def _is_retryable_llm_error(exc: Exception) -> bool:
    """判断异常是否属于 auth/connection/timeout 类(可熔断降级)。"""
    if isinstance(exc, TimeoutError):
        return True
    name = type(exc).__name__
    try:
        import openai
        if isinstance(exc, (
            openai.APIConnectionError, openai.APITimeoutError,
            openai.AuthenticationError, openai.PermissionDeniedError,
            openai.RateLimitError, openai.InternalServerError,
        )):
            return True
    except Exception:
        pass
    name_l = name.lower()
    _circuit_keys = ("auth", "timeout", "connection", "connect", "rate", "permission", "apierror")
    if any(k in name_l for k in _circuit_keys):
        return True
    msg = str(exc).lower()
    if any(k in msg for k in (
        "401", "403", "429", "502", "503", "504",
        "timeout", "timed out", "unauthorized", "connection",
        "refused", "reset by peer", "api key", "apikey", "invalid_api_key",
    )):
        return True
    return False


def _invoke_one(llm, messages, timeout_sec: int) -> str:
    """对单个 LLM 实例执行调用(带超时)。"""
    # 主线程内判断模型类型, 避免把 langchain 导入成本计入调用超时
    chat_model = _is_chat_model(llm)

    if chat_model:
        def _call():
            response = llm.invoke(messages)
            return response.content
    else:
        prompt = "\n".join(
            f"{getattr(m, 'type', 'user')}: {m.content}" for m in messages
        )
        def _call():
            response = llm.invoke(prompt)
            return response if isinstance(response, str) else str(response)

    pool = _get_llm_executor()
    future = pool.submit(_call)
    try:
        return future.result(timeout=timeout_sec)
    except concurrent.futures.TimeoutError:
        raise TimeoutError(f"LLM 调用超时 ({timeout_sec}s)") from None


def _record_llm_metric(succeeded: bool, start_time: float) -> None:
    """LLM 调用指标埋点 — 任何异常仅记日志, 不影响降级链/业务。"""
    try:
        from bobanana.observability import metrics
        if succeeded:
            metrics.inc("llm_calls_total")
            metrics.observe("llm_call_seconds", time.monotonic() - start_time)
        else:
            metrics.inc("llm_errors_total")
    except Exception as e:  # noqa: BLE001
        logger.debug("LLM 指标埋点失败: %s", e)


def llm_invoke(system_prompt: str, user_prompt: str, timeout_sec: int = None) -> str:
    """调用 LLM，返回文本结果。

    按降级链顺序尝试 provider;auth/connection/timeout 类失败会熔断该
    provider(冷却 60s)并尝试下一个, 全部失败抛 SWError(SW-LLM-503)。
    """
    from langchain_core.messages import HumanMessage, SystemMessage

    from bobanana.config import LLM_TIMEOUT_SEC
    from bobanana.errors import SW_LLM_503, SWError

    start_time = time.monotonic()
    succeeded = False
    try:
        if timeout_sec is None:
            timeout_sec = LLM_TIMEOUT_SEC

        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=user_prompt),
        ]

        last_err = None
        tried = []
        for provider in _providers():
            if _is_circuit_open(provider):
                continue
            llm = get_llm(provider)
            if llm is None:
                continue
            tried.append(provider)
            try:
                result = _invoke_one(llm, messages, timeout_sec)
                succeeded = True
                return result
            except Exception as e:
                if _is_retryable_llm_error(e):
                    _trip_circuit(provider)
                    last_err = e
                    logger.warning("LLM provider %s 调用失败, 熔断 %.0fs: %s",
                                   provider, _PROVIDER_COOLDOWN_SEC, e)
                else:
                    logger.error("LLM 调用失败(非可重试): %s", e)
                    raise

        detail = f"providers={','.join(tried) or 'none'}; last_error={last_err}"
        raise SWError(error_code=SW_LLM_503, message="LLM 服务暂时不可用，请稍后重试", detail=detail)
    finally:
        _record_llm_metric(succeeded, start_time)


def llm_stream(system_prompt: str, user_prompt: str):
    """流式生成器, 逐块 yield 文本。失败时降级为一次 llm_invoke 后单块 yield。"""
    from langchain_core.messages import HumanMessage, SystemMessage

    messages = [
        SystemMessage(content=system_prompt),
        HumanMessage(content=user_prompt),
    ]
    try:
        llm = get_llm()
        if llm is None or not hasattr(llm, "stream") or not _is_chat_model(llm):
            raise RuntimeError("当前 provider 不支持流式, 降级")
        for chunk in llm.stream(messages):
            content = getattr(chunk, "content", "")
            if isinstance(content, str) and content:
                yield content
            elif isinstance(content, list):
                for part in content:
                    if isinstance(part, str):
                        yield part
                    elif isinstance(part, dict) and part.get("type") == "text":
                        yield part.get("text", "")
    except Exception as e:
        logger.warning("LLM 流式失败, 降级为单次 llm_invoke: %s", e)
        try:
            text = llm_invoke(system_prompt, user_prompt)
            if text:
                yield text
        except Exception:
            yield ""
