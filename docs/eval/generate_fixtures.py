"""生成 Agent/检索评测用的 10 份小型文档夹具(4 MD + 3 PDF + 3 Word)。

用法:
    .venv-linux/bin/python docs/eval/generate_fixtures.py

产物写入 docs/eval/fixtures/,可重复运行覆盖。
内容为课程风格知识点,与 docs/eval/agent_instructions.jsonl、
docs/eval/retrieval_queries.jsonl 配套。
"""
from __future__ import annotations

from pathlib import Path

HERE = Path(__file__).resolve().parent
OUT = HERE / "fixtures"

# ── 文档内容: {输出文件名: (文档标题, [(小节标题, [段落, ...]), ...])} ──
MD_DOCS: dict[str, tuple[str, list[tuple[str, list[str]]]]] = {
    "md_digital_logic_gates.md": (
        "数字逻辑: 逻辑门基础",
        [
            ("与门(AND)", [
                "与门是最基本的逻辑门之一,只有当所有输入都为高电平(逻辑1)时,输出才为高电平。与门的逻辑表达式为 Y = A · B,读作 Y 等于 A 与 B。两输入与门的真值表共有四种组合:00、01、10、11,只有输入 11 时输出为 1,其余组合输出均为 0。",
                "与门可以用生活中的例子理解:一扇需要两把钥匙同时转动才能打开的门,只有两把钥匙都在位时门才打开。在数字电路中,与门常用于条件判断,例如只有使能信号和数据信号同时有效时才允许数据通过。",
            ]),
            ("或门(OR)", [
                "或门的规则是:只要任意一个输入为高电平,输出就为高电平。逻辑表达式为 Y = A + B。或门真值表中,除输入 00 外,其余组合(01、10、11)输出都是 1。",
                "生活中或门的例子是并联开关控制的楼道灯:任何一个开关闭合都能点亮。电路设计中,或门常用于多路条件触发,如多个中断源中任意一个有效就产生中断请求。",
            ]),
            ("非门与异或门", [
                "非门(NOT)只有一个输入,输出与输入相反,表达式 Y = A'(A 非)。异或门(XOR)在两输入不相同时输出 1,相同时输出 0,表达式 Y = A ⊕ B,常用于奇偶校验和加法器的本位和计算。",
                "与非门(NAND)是与门加非门,被称为万能门:任何逻辑电路都可以只用与非门搭建,这是数字电路集成化的基础。",
            ]),
        ],
    ),
    "md_digital_logic_combinational.md": (
        "数字逻辑: 组合逻辑电路",
        [
            ("半加器与全加器", [
                "半加器实现两个一位二进制数的相加,输出本位和 S 与进位 C:S = A ⊕ B,C = A · B。半加器没有考虑低位进位,只能做最低位加法。",
                "全加器在半加器基础上增加进位输入 Cin,实现 A、B、Cin 三个数的求和:S = A ⊕ B ⊕ Cin,Cout = AB + (A ⊕ B)Cin。多个全加器级联可构成多位加法器,是 CPU 算术逻辑单元的核心。",
            ]),
            ("编码器与译码器", [
                "编码器把 2^n 个输入信号编码为 n 位二进制码,例如 8-3 编码器把 8 个按键编码为 3 位码。译码器是编码器的逆过程,把 n 位二进制码转换为 2^n 个输出中的一路有效,如 3-8 译码器,常用于地址译码和七段数码管驱动。",
            ]),
            ("多路选择器", [
                "多路选择器(MUX)根据选择信号从多个输入中挑选一路输出,例如 4 选 1 多路选择器用 2 位选择信号控制 4 路数据。MUX 是实现组合逻辑函数的通用元件,也可作为数据通路中的开关。",
            ]),
        ],
    ),
    "md_kirchhoff.md": (
        "电路分析: 基尔霍夫定律",
        [
            ("基尔霍夫电流定律(KCL)", [
                "KCL 指出:任意时刻,流入电路中任一节点的电流之和等于流出该节点的电流之和,本质是电荷守恒。公式写作 ΣI入 = ΣI出。列写节点方程时,可约定流入为正、流出为负,所有支路电流代数和为零。",
                "KCL 适用于任何集总参数电路,包括非线性电路。典型应用是把复杂电路中的节点电流关系转化为线性方程组求解。",
            ]),
            ("基尔霍夫电压定律(KVL)", [
                "KVL 指出:沿任意闭合回路绕行一周,各元件电压降的代数和为零,本质是能量守恒。列方程时先标定绕行方向,电压升取正、电压降取负。",
                "KVL 与欧姆定律结合,可对每个独立回路列出方程。节点电压法与回路电流法都是基于基尔霍夫定律的系统化求解方法:节点电压法以节点电位为未知量,回路电流法以假想回路电流为未知量,两者列方程数量少、适合手工计算。",
            ]),
        ],
    ),
    "md_thevenin.md": (
        "电路分析: 戴维南定理与诺顿定理",
        [
            ("戴维南定理", [
                "戴维南定理指出:任何线性含源二端网络,对外电路而言都可以等效为一个电压源与电阻串联的电路,电压源电压等于二端网络的开路电压 Uoc,串联电阻等于网络内部独立源置零后从端口看入的等效电阻 Req。",
                "求等效电阻的方法:将网络内所有独立电压源短路、独立电流源开路,再从端口计算电阻;若电路含有受控源,则采用外加电源法(端口加电压源,求端口电流,两者之比即 Req)。",
            ]),
            ("诺顿定理与最大功率传输", [
                "诺顿定理是戴维南定理的对偶形式:线性含源二端网络可等效为电流源与电阻并联,电流源电流等于端口短路电流 Isc。戴维南与诺顿等效可互相转换:Uoc = Isc × Req。",
                "最大功率传输定理:当负载电阻 RL 等于等效内阻 Req 时,负载获得最大功率 Pmax = Uoc²/(4Req)。该定理用于匹配电路设计,如音频功放与扬声器阻抗匹配。",
            ]),
        ],
    ),
}

PDF_DOCS: dict[str, tuple[str, list[tuple[str, list[str]]]]] = {
    "pdf_diode.pdf": (
        "模拟电子: 半导体二极管",
        [
            ("PN 结", [
                "P 型半导体多子为空穴,N 型半导体多子为电子。P 区与 N 区接触后,交界处多子扩散复合,留下不能移动的离子,形成空间电荷区(耗尽层),内部建立起由 N 指向 P 的内建电场。",
                "PN 结具有单向导电性:外加正向电压(正偏)时,耗尽层变窄,电流随电压指数增长;外加反向电压(反偏)时,耗尽层变宽,仅有极小的反向饱和电流。",
            ]),
            ("二极管伏安特性与稳压管", [
                "二极管的伏安特性:正向导通需超过死区电压(硅管约 0.7V),导通后电压基本不变而电流迅速增大;反向电压超过击穿电压时发生击穿,普通二极管此时会损坏。",
                "稳压二极管(齐纳管)工作在反向击穿区:在击穿状态下端电压几乎不随电流变化,因此可提供稳定基准电压。使用稳压管必须串联限流电阻,防止电流过大烧毁。",
            ]),
        ],
    ),
    "pdf_transistor.pdf": (
        "模拟电子: 三极管放大电路",
        [
            ("共射极放大电路", [
                "三极管(BJT)有三个电极:基极 B、集电极 C、发射极 E,分为 NPN 与 PNP 两类。共射极放大电路以发射极为公共端,信号从基极输入、集电极输出,兼具电压放大与电流放大能力。",
                "放大条件:发射结正偏、集电结反偏。静态工作点由偏置电阻决定,设置合适的静态工作点是保证不失真放大的前提。",
            ]),
            ("放大参数与失真", [
                "电流放大倍数 β = Ic / Ib,是衡量三极管放大能力的重要参数。电压放大倍数 Au = -β × Rc / rbe,负号表示共射极电路输出与输入反相。",
                "静态工作点过高会产生饱和失真(输出波形底部削平),过低则产生截止失真(顶部削平)。合理选择偏置电路(如分压式偏置)可稳定静态工作点,减小温度对放大性能的影响。",
            ]),
        ],
    ),
    "pdf_maxwell.pdf": (
        "电磁学: 麦克斯韦方程组",
        [
            ("四个基本方程", [
                "麦克斯韦方程组由四个方程构成:高斯定律(电场的有源特性,电荷是电场的源)、磁高斯定律(磁场无源,磁单极子不存在)、法拉第电磁感应定律(变化的磁场产生涡旋电场)、安培-麦克斯韦定律(电流和变化的电场都产生磁场)。",
            ]),
            ("位移电流与电磁波", [
                "麦克斯韦在安培环路定律中引入位移电流概念:变化的电场等价于一种电流,也能激发磁场。位移电流的引入使方程组对称完整,从而在理论上预言了电磁波的存在,并推出电磁波在真空中以光速传播。",
                "麦克斯韦方程组是经典电磁学的集大成者:它统一了电、磁、光三种现象。赫兹实验验证了电磁波,为无线通信奠定了理论基础。",
            ]),
        ],
    ),
}

DOCX_DOCS: dict[str, tuple[str, list[tuple[str, list[str]]]]] = {
    "docx_von_neumann.docx": (
        "计算机组成: 冯·诺依曼结构",
        [
            ("五大部件", [
                "冯·诺依曼计算机由五大部件组成:运算器、控制器、存储器、输入设备、输出设备。运算器完成算术与逻辑运算,控制器负责取指令、译码并指挥各部件协调工作,二者合称中央处理器 CPU。",
            ]),
            ("存储程序与指令周期", [
                "存储程序是冯·诺依曼结构的核心思想:程序与数据都以二进制形式存放在同一个存储器中,机器按地址依次取出指令执行,程序可以通过修改存储内容实现跳转。",
                "指令周期包括取指、译码、执行三个阶段。程序计数器 PC 保存下一条指令的地址,每取完一条指令自动加一;遇到跳转指令时,PC 被置为目标地址。",
            ]),
        ],
    ),
    "docx_stack_queue.docx": (
        "数据结构: 栈与队列",
        [
            ("栈", [
                "栈是受限的线性表,只允许在表的一端(栈顶)进行插入和删除,遵循后进先出(LIFO)原则。基本操作有入栈 push、出栈 pop、读栈顶 peek。",
                "栈的典型应用:函数调用时保存返回地址与局部变量(调用栈)、表达式求值与括号匹配、深度优先搜索、浏览器的后退历史。",
            ]),
            ("队列", [
                "队列也只允许在一端(队尾)插入、另一端(队头)删除,遵循先进先出(FIFO)原则。基本操作有入队 enqueue、出队 dequeue。",
                "顺序存储的队列存在假溢出问题,常用循环队列解决:队头队尾指针到达数组末尾后回到开头。队列的应用:任务调度、广度优先搜索、消息缓冲、打印任务排队。",
            ]),
        ],
    ),
    "docx_process_thread.docx": (
        "操作系统: 进程与线程",
        [
            ("进程", [
                "进程是程序的一次执行,是资源分配的基本单位。每个进程拥有独立的地址空间,由进程控制块 PCB 记录进程标识、状态、程序计数器、寄存器现场、内存指针等管理信息。",
                "进程状态的转换:就绪、运行、阻塞。CPU 调度在就绪队列中挑选进程运行;进程等待 I/O 时进入阻塞态,完成后回到就绪态。",
            ]),
            ("线程与并发", [
                "线程是 CPU 调度的基本单位,同一进程内的线程共享该进程的地址空间、打开的文件等资源,但各自拥有独立的栈和寄存器现场,因此线程创建与切换开销远小于进程。",
                "并发指多个任务在一段时间内交替执行(单核时间片轮转),并行指多个任务在同一时刻真正同时执行(多核)。上下文切换是 CPU 从执行一个进程/线程切换到另一个时保存和恢复现场的过程,切换过频会带来性能开销。",
            ]),
        ],
    ),
}


def _write_md(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    lines = [f"# {title}", ""]
    for sec_title, paras in sections:
        lines.append(f"## {sec_title}")
        lines.append("")
        for p in paras:
            lines.append(p)
            lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def _write_pdf(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    import fitz  # PyMuPDF

    doc = fitz.open()
    page = doc.new_page()
    rect = fitz.Rect(72, 72, 523, 770)  # A4 页边距
    y = rect.y0

    def emit(text: str, size: float, bold: bool = False) -> None:
        nonlocal y, page
        while text:
            # 预留行高
            avail = rect.y1 - y
            if avail < size * 1.5:
                page = doc.new_page()
                y = rect.y0
            rc = page.insert_textbox(
                fitz.Rect(rect.x0, y, rect.x1, rect.y1),
                text,
                fontsize=size,
                fontname="china-s",
                lineheight=size * 1.45,
            )
            if rc >= 0:
                y += rc
                break
            # 放不下:按估算折半切分
            est = max(1, int(len(text) * (rect.y1 - y) / (size * 1.45 * (rect.x1 - rect.x0) / 2)))
            chunk, rest = text[:est], text[est:]
            rc2 = page.insert_textbox(
                fitz.Rect(rect.x0, y, rect.x1, rect.y1), chunk,
                fontsize=size, fontname="china-s", lineheight=size * 1.45,
            )
            y += max(rc2, size * 1.45)
            text = rest

    emit(title, 16)
    y += 8
    for sec_title, paras in sections:
        emit(f"■ {sec_title}", 12.5)
        y += 3
        for p in paras:
            emit(p, 10.5)
        y += 6
    doc.save(str(path))
    doc.close()


def _write_docx(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    import docx

    d = docx.Document()
    d.add_heading(title, level=0)
    for sec_title, paras in sections:
        d.add_heading(sec_title, level=1)
        for p in paras:
            d.add_paragraph(p)
    d.save(str(path))


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for name, (title, sections) in MD_DOCS.items():
        _write_md(OUT / name, title, sections)
    for name, (title, sections) in PDF_DOCS.items():
        _write_pdf(OUT / name, title, sections)
    for name, (title, sections) in DOCX_DOCS.items():
        _write_docx(OUT / name, title, sections)
    print(f"已生成 {len(MD_DOCS) + len(PDF_DOCS) + len(DOCX_DOCS)} 份夹具到 {OUT}")


if __name__ == "__main__":
    main()
